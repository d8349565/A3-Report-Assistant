"""
A3 Report Assistant – Word Edition v1.0.0
========================================
* 完整的管理员权限控制和配置管理系统
* 支持AJAX无刷新文件下载和生成进度提示
* 防重复点击机制和任务去重保护
* 基于DeepSeek AI的智能A3报告生成
* 依赖：flask python-docx openai waitress
"""

from __future__ import annotations
import datetime as _dt
import json
import re
import sys
import os
import webbrowser
from pathlib import Path
from typing import Dict, List

from flask import (
    Flask,
    request,
    render_template,
    render_template_string,
    send_file,
    redirect,
    url_for,
    flash,
    jsonify,
    session,
)
import openai
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

import config

def ensure_api_key():
    """检查API Key是否存在，Docker环境下跳过交互式输入"""
    api_key = os.getenv("DEEPSEEK_API_KEY") or getattr(config, "DEEPSEEK_API_KEY", None)
    if not api_key or not api_key.strip():
        # 在Docker环境中，如果没有API Key则退出
        if os.getenv("FLASK_ENV") == "production":
            print("错误: 未设置 DEEPSEEK_API_KEY 环境变量")
            print("请在 .env 文件中设置 DEEPSEEK_API_KEY")
            sys.exit(1)
        else:
            # 开发环境下保持原有交互式输入
            key = input("请输入 DeepSeek API Key：").strip()
            # 写入 config.py
            with open("config.py", "r", encoding="utf-8") as f:
                lines = f.readlines()
            with open("config.py", "w", encoding="utf-8") as f:
                for line in lines:
                    if line.strip().startswith("DEEPSEEK_API_KEY"):
                        f.write(f'DEEPSEEK_API_KEY = "{key}"\n')
                    else:
                        f.write(line)
            print("API Key 已保存，请重新启动程序。")
            sys.exit(0)

# -------------------------------------------------------------
# Configuration
# -------------------------------------------------------------
OUTPUT_DIR = Path(__file__).parent / config.OUTPUT_DIR_NAME
OUTPUT_DIR.mkdir(exist_ok=True)

# 从config导入配置
GUIDE: List[Dict[str, str]] = config.GUIDE
# 优先从环境变量读取 API Key，用于Docker部署
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY") or config.DEEPSEEK_API_KEY

GUIDE_MAP = {g["id"]: g for g in GUIDE}

app = Flask(__name__, template_folder='templates')
app.secret_key = config.APP_SECRET_KEY

# 全局变量：存储正在生成报告的标识
generating_reports = set()

# -------------------------------------------------------------
# Helper functions
# -------------------------------------------------------------

def require_access(f):
    """访问权限验证装饰器"""
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # 检查是否已通过访问验证
        if not session.get('access_granted'):
            return redirect(url_for('access_login'))
        return f(*args, **kwargs)
    return decorated_function

def reload_config():
    """重新加载配置文件"""
    import importlib
    importlib.reload(config)
    global GUIDE, GUIDE_MAP, DEEPSEEK_API_KEY
    GUIDE = config.GUIDE
    GUIDE_MAP = {g["id"]: g for g in GUIDE}
    DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY") or config.DEEPSEEK_API_KEY

def save_config(new_config):
    """保存配置到文件"""
    try:
        with open("config.py", "r", encoding="utf-8") as f:
            content = f.read()
        
        # 更新基本配置
        content = re.sub(
            r'DEEPSEEK_API_KEY = ".*?"',
            f'DEEPSEEK_API_KEY = "{new_config.get("deepseek_api_key", "")}"',
            content
        )
        content = re.sub(
            r'DEEPSEEK_BASE_URL = ".*?"',
            f'DEEPSEEK_BASE_URL = "{new_config.get("deepseek_base_url", "")}"',
            content
        )
        content = re.sub(
            r'MODEL_NAME = ".*?"',
            f'MODEL_NAME = "{new_config.get("model_name", "")}"',
            content
        )
        content = re.sub(
            r'DOC_FONT_NAME = ".*?"',
            f'DOC_FONT_NAME = "{new_config.get("doc_font_name", "")}"',
            content
        )
        content = re.sub(
            r'DOC_TITLE_TEMPLATE = ".*?"',
            f'DOC_TITLE_TEMPLATE = "{new_config.get("doc_title_template", "")}"',
            content
        )
        content = re.sub(
            r'WEB_ACCESS_PASSWORD = ".*?"',
            f'WEB_ACCESS_PASSWORD = "{new_config.get("web_access_password", "")}"',
            content
        )
        
        # 更新系统提示词
        for key in ["default", "step_guidance", "validation", "optimization"]:
            pattern = f'"{key}": """.*?"""'
            prompt_key = f"system_prompt_{key}"
            prompt_content = new_config.get(prompt_key, "").replace('"', '\\"')
            replacement = f'"{key}": """{prompt_content}"""'
            content = re.sub(pattern, replacement, content, flags=re.DOTALL)
        
        # 更新A3步骤配置
        step_ids = new_config.getlist("step_id[]") if "step_id[]" in new_config else []
        if step_ids:
            step_titles = new_config.getlist("step_title[]")
            step_purposes = new_config.getlist("step_purpose[]")
            step_tools = new_config.getlist("step_tools[]")
            step_focuses = new_config.getlist("step_focus[]")
            
            guide_content = "GUIDE = [\n"
            for i in range(len(step_ids)):
                focus_escaped = step_focuses[i].replace('"', '\\"')
                guide_content += f'''    {{
        "id": "{step_ids[i]}",
        "title": "{step_titles[i]}",
        "purpose": "{step_purposes[i]}",
        "tools": "{step_tools[i]}",
        "focus": "{focus_escaped}",
    }},\n'''
            guide_content += "]"
            
            content = re.sub(
                r'GUIDE = \[.*?\]',
                guide_content,
                content,
                flags=re.DOTALL
            )
        
        with open("config.py", "w", encoding="utf-8") as f:
            f.write(content)
        
        return True
    except Exception as e:
        print(f"保存配置失败: {e}")
        return False

def sanitize_filename(text: str) -> str:
    return re.sub(r"[^\w\- ]", "", text).strip()[:40] or "A3Report"


def next_docx_path(topic: str) -> Path:
    base = f"{_dt.datetime.now():%Y%m%d_%H%M}_{sanitize_filename(topic)}"
    idx = 1
    while True:
        p = OUTPUT_DIR / f"{base}_v{idx}.docx"
        if not p.exists():
            return p
        idx += 1


def call_deepseek(prompt: str, api_key: str = None, step_id: str = None) -> str:
    api_key = api_key or DEEPSEEK_API_KEY
    # 如果有step_id，拼接system prompt
    if step_id and step_id in GUIDE_MAP:
        st = GUIDE_MAP[step_id]
        sys_prompt = config.SYSTEM_PROMPTS["step_guidance"].format(
            title=st['title'],
            purpose=st['purpose'],
            tools=st['tools'],
            focus=st['focus']
        )
        messages = [
            {"role": "system", "content": sys_prompt},
            {"role": "user", "content": prompt},
        ]
        return call_deepseek_multi(messages, api_key)
    # 否则用默认system prompt
    client = openai.OpenAI(api_key=api_key, base_url=config.DEEPSEEK_BASE_URL)
    try:
        resp = client.chat.completions.create(
            model=config.MODEL_NAME,
            messages=[
                {"role": "system", "content": config.SYSTEM_PROMPTS["default"]},
                {"role": "user", "content": prompt},
            ],
            stream=False,
        )
        return resp.choices[0].message.content.strip()
    except Exception as exc:
        return f"<LLM 调用失败> {exc}"


def build_doc(topic: str, user_inputs: Dict[str, str], suggestions: Dict[str, str]) -> Path:
    doc = Document()
    h = doc.add_heading(config.DOC_TITLE_TEMPLATE.format(topic=topic), level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in h.runs:
        set_font_simsun(run)

    for idx, st in enumerate(GUIDE, 1):
        heading = doc.add_heading(f"Step {idx}: {st['title']}", level=2)
        for run in heading.runs:
            set_font_simsun(run)
        p1 = doc.add_paragraph(user_inputs.get(st["id"], "(用户未填写)"))
        for run in p1.runs:
            set_font_simsun(run)
        p2 = doc.add_paragraph("优化建议：", style="Intense Quote")
        for run in p2.runs:
            set_font_simsun(run)
        p3 = doc.add_paragraph(suggestions.get(st["id"], "-"))
        for run in p3.runs:
            set_font_simsun(run)

    path = next_docx_path(topic)
    doc.save(path)
    return path

# -------------------------------------------------------------
# Web templates
# -------------------------------------------------------------

# 现在使用外部模板文件，不再需要内嵌HTML

# -------------------------------------------------------------
# Routes
# -------------------------------------------------------------

@app.route("/")
@require_access
def index():
    step_ids_json = json.dumps([g["id"] for g in GUIDE])
    return render_template('index.html', guide=GUIDE, step_ids=step_ids_json)


@app.route("/access")
def access_login():
    """网页访问密码登录页面"""
    return render_template('access_login.html')


@app.route("/access", methods=["POST"])
def access_login_post():
    """处理网页访问密码验证"""
    password = request.form.get('password', '')
    if password == config.WEB_ACCESS_PASSWORD:
        session['access_granted'] = True
        return redirect(url_for('index'))
    else:
        flash("访问密码错误，请重试。")
        return redirect(url_for('access_login'))


@app.route("/access/logout")
def access_logout():
    """退出访问权限"""
    session.pop('access_granted', None)
    session.pop('admin_logged_in', None)  # 同时清除管理员登录状态
    flash("已安全退出。")
    return redirect(url_for('access_login'))


@app.route("/admin")
@require_access
def admin():
    """管理员配置页面"""
    # 检查是否已登录
    if not session.get('admin_logged_in'):
        return redirect(url_for('admin_login'))
    return render_template('admin.html', config=config)


@app.route("/admin/login")
@require_access
def admin_login():
    """管理员登录页面"""
    return render_template('admin_login.html')


@app.route("/admin/login", methods=["POST"])
def admin_login_post():
    """处理管理员登录"""
    password = request.form.get('password', '')
    if password == config.ADMIN_PASSWORD:
        session['admin_logged_in'] = True
        flash("登录成功！")
        return redirect(url_for('admin'))
    else:
        flash("密码错误，请重试。")
        return redirect(url_for('admin_login'))


@app.route("/admin/logout")
def admin_logout():
    """管理员登出"""
    session.pop('admin_logged_in', None)
    flash("已安全登出。")
    return redirect(url_for('index'))


@app.route("/admin", methods=["POST"])
@require_access
def admin_save():
    """保存管理员配置"""
    # 检查是否已登录
    if not session.get('admin_logged_in'):
        return redirect(url_for('admin_login'))
    
    try:
        if save_config(request.form):
            reload_config()
            flash("配置保存成功！")
        else:
            flash("配置保存失败，请检查输入内容。")
    except Exception as e:
        flash(f"保存配置时发生错误：{str(e)}")
    
    return redirect(url_for('admin'))


@app.route("/validate", methods=["POST"])
@require_access
def validate():
    data = request.get_json(force=True)
    step_id = data.get("step_id")
    inputs = data.get("inputs", {})
    history = data.get("history", [])  # 新增：历史对话
    if step_id not in GUIDE_MAP:
        return jsonify({"error": "参数错误"}), 400

    context = "\n".join([f"{GUIDE_MAP[k]['title']}: {v}" for k, v in inputs.items() if v.strip()])
    # 拼接当前步骤的A3指导信息
    st = GUIDE_MAP[step_id]
    sys_prompt = config.SYSTEM_PROMPTS["step_guidance"].format(
        title=st['title'],
        purpose=st['purpose'],
        tools=st['tools'],
        focus=st['focus']
    )
    user_prompt = config.SYSTEM_PROMPTS["validation"].format(
        context=context,
        title=st['title']
    )
    messages = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt}
    ]
    for msg in history:
        if msg.get("role") in ("user","assistant") and msg.get("content"):
            messages.append({"role": msg["role"], "content": msg["content"]})
    suggestion = call_deepseek_multi(messages)
    return jsonify({"suggestion": suggestion})


@app.route("/generate", methods=["POST"])
@require_access
def generate():
    # 生成唯一的任务ID
    import hashlib
    import time
    task_data = f"{request.remote_addr}_{time.time()}"
    task_id = hashlib.md5(task_data.encode()).hexdigest()[:12]
    
    # 检查是否已有任务在进行
    if task_id in generating_reports:
        return jsonify({"error": "报告生成中，请稍候..."}), 429
    
    # 标记任务开始
    generating_reports.add(task_id)
    
    try:
        user_inputs = {g["id"]: request.form.get(g["id"], "").strip() for g in GUIDE}
        topic = user_inputs.get("step1", "A3_Topic")[:30]

        suggestions: Dict[str, str] = {}
        for st in GUIDE:
            content = user_inputs.get(st["id"], "")
            if not content:
                suggestions[st["id"]] = "(用户未填写)"
                continue
            prompt = config.SYSTEM_PROMPTS["optimization"].format(
                title=st['title'],
                content=content
            )
            suggestions[st["id"]] = call_deepseek(prompt, step_id=st["id"])

        path = build_doc(topic, user_inputs, suggestions)
        
        # 任务完成，移除标记
        generating_reports.discard(task_id)
        
        return send_file(path, as_attachment=True)
    
    except Exception as e:
        # 发生错误时也要移除标记
        generating_reports.discard(task_id)
        flash(f"生成报告时发生错误：{str(e)}")
        return redirect(url_for('index'))


@app.route("/generate/status/<task_id>")
@require_access
def generate_status(task_id):
    """检查生成任务状态"""
    is_generating = task_id in generating_reports
    return jsonify({"generating": is_generating})

# -------------------------------------------------------------
# 新增多轮对话支持

def call_deepseek_multi(messages, api_key: str = None) -> str:
    api_key = api_key or DEEPSEEK_API_KEY
    client = openai.OpenAI(api_key=api_key, base_url=config.DEEPSEEK_BASE_URL)
    try:
        resp = client.chat.completions.create(
            model=config.MODEL_NAME,
            messages=messages,
            stream=False,
        )
        return resp.choices[0].message.content.strip()
    except Exception as exc:
        return f"<LLM 调用失败> {exc}"

def set_font_simsun(run):
    run.font.name = config.DOC_FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), config.DOC_FONT_NAME)

# -------------------------------------------------------------
if __name__ == "__main__":
    ensure_api_key()
    webbrowser.open("http://localhost:9998")
    if getattr(sys, "frozen", False):
        from waitress import serve
        serve(app, host="0.0.0.0", port=9998)
    else:
        app.run(host="0.0.0.0", port=9998, debug=False)
